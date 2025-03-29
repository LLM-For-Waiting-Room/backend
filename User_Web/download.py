import json
from docx import Document

with open('/home/sunanhe/qiboyuan/MedicalWeb/User_Web/data.json', 'r', encoding='utf-8') as file:
    data = json.load(file)


data_list = data[0]['rows']
# 取最后十个数据
last_ten_items = data_list[-10:]

for index,item in enumerate(last_ten_items):
    doc = Document()

    conv_list = json.loads(item[4])
    for i in conv_list:
        doc.add_paragraph(str(i))
    doc.add_paragraph('\n')
    doc.add_paragraph(str(item[7]))

    print(item[4],'\n')
    print(item[7],'\n')
    doc.save(f"test{index}.docx")